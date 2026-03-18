"""Generate a modern-styled Word document from the Executive Brief."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

DIR = os.path.dirname(os.path.abspath(__file__))

# Modern color palette
DARK = RGBColor(0x1A, 0x1A, 0x2E)       # Near-black for body text
ACCENT = RGBColor(0x4F, 0x46, 0xE5)      # Indigo-600 (headings)
ACCENT_LIGHT = "6366F1"                    # Indigo-500 hex for XML
MUTED = RGBColor(0x64, 0x74, 0x8B)        # Slate-500 (secondary text)
TABLE_HEADER_BG = "4F46E5"                 # Indigo for table headers
TABLE_ALT_BG = "F8FAFC"                    # Slate-50 for alternating rows
RULE_COLOR = "E2E8F0"                      # Slate-200 for dividers
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def remove_cell_borders(table):
    """Set thin subtle borders on table."""
    tbl = table._tbl
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:color="{RULE_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="4" w:color="{RULE_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:color="{RULE_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="4" w:color="{RULE_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:color="{RULE_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:color="{RULE_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tbl.tblPr.append(borders)


def style_table(table, col_widths=None):
    """Apply modern styling to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_cell_borders(table)

    # Header row
    for cell in table.rows[0].cells:
        set_cell_shading(cell, TABLE_HEADER_BG)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.name = "Calibri"

    # Data rows
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            if i % 2 == 0:
                set_cell_shading(cell, TABLE_ALT_BG)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = DARK
                    run.font.name = "Calibri"

    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(w)


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:color="{RULE_COLOR}" w:space="1"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(26)
    run.font.color.rgb = ACCENT
    run.font.bold = True


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = DARK
    run.font.bold = True


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = ACCENT
    run.font.bold = True


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    return p


def body_rich(doc, parts):
    """Add paragraph with mixed bold/normal runs. parts = [(text, bold), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(18)
    for text, bold in parts:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK
        run.bold = bold
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK
        r.bold = True
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK
    else:
        p.clear()
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    # Light background via shading
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
    pPr.append(shading)
    run = p.add_run(text)
    run.font.name = "Cascadia Code"
    run.font.size = Pt(9)
    run.font.color.rgb = DARK


def caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    run.italic = True


def build():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = DARK

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── TITLE PAGE ──
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HEALTH TRACKER")
    run.font.name = "Calibri"
    run.font.size = Pt(36)
    run.font.color.rgb = ACCENT
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Executive Brief")
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("Personal Health Analytics Platform")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("AI-Architected  \u00b7  Evidence-Based  \u00b7  Fully Automated")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT

    doc.add_page_break()

    # ── WHAT IT IS ──
    h1(doc, "What It Is")

    body(doc,
        "Health Tracker is a personal health analytics platform that automatically collects "
        "50+ daily biometrics from a Garmin wearable, combines them with manual habit tracking, "
        "and produces a daily Readiness Score \u2014 a single number (1\u201310) that tells you how "
        "prepared your body and mind are for the day ahead."
    )

    body(doc,
        "Every morning, a push notification arrives on your phone with your score, a sleep summary, "
        "training load status, and 3\u20135 actionable recommendations grounded in peer-reviewed research. "
        "No dashboards to check, no apps to open. The insight comes to you."
    )

    body(doc,
        "The system was architected and engineered entirely with Claude (Anthropic) serving as both "
        "the development partner and the runtime analytical engine \u2014 a demonstration of what "
        "AI-augmented personal health infrastructure looks like when built with rigor."
    )

    add_divider(doc)

    # ── THE PROBLEM ──
    h1(doc, "The Problem")

    body(doc,
        "Wearables collect enormous amounts of data. Garmin alone produces 50+ metrics per day across "
        "sleep, heart rate variability, training load, stress, and recovery. But the native apps present "
        "this data in isolation \u2014 sleep in one screen, HRV in another, training in a third. "
        "No synthesis. No context. No personalization for medical conditions."
    )

    body(doc,
        "The result: most people glance at a number, don\u2019t know what it means relative to their "
        "baseline, and move on. The data exists but never becomes insight."
    )

    body(doc,
        "Health Tracker solves this by building a personal analytical layer on top of raw wearable "
        "data \u2014 one that understands your physiology, tracks your research interests, and delivers "
        "concise, evidence-backed guidance every morning."
    )

    add_divider(doc)

    # ── HOW IT WORKS ──
    h1(doc, "How It Works")
    h2(doc, "Architecture Overview")

    # Architecture diagram as styled table
    arch_table = doc.add_table(rows=4, cols=4)
    arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_cell_borders(arch_table)
    headers = ["Ingestion", "Storage", "Analysis", "Delivery"]
    items = [
        ["Garmin API\n(50+ metrics)", "Google Sheets\n(source of truth)", "Readiness Engine\n(sigmoid z-scores)", "Push Notification\n(Pushover API)"],
        ["Voice Logger\n(Claude Haiku NLP)", "SQLite Mirror\n(offline backup)", "Knowledge Base\n(100+ entries, 8 domains)", "Color-Graded Sheets\n(research thresholds)"],
        ["Garmin Export\n(historical bulk)", "Raw Data Archive", "Health Profile\n(conditions, biomarkers)", "Weekly Validation\n(predictions vs outcomes)"],
    ]
    for i, h in enumerate(headers):
        cell = arch_table.rows[0].cells[i]
        set_cell_shading(cell, TABLE_HEADER_BG)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.bold = True

    for r, row_data in enumerate(items):
        for c, val in enumerate(row_data):
            cell = arch_table.rows[r + 1].cells[c]
            if (r + 1) % 2 == 0:
                set_cell_shading(cell, TABLE_ALT_BG)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            run.font.color.rgb = DARK

    h2(doc, "Data Pipeline")

    body(doc,
        "A Python script runs nightly at 8 PM via the operating system\u2019s native scheduler "
        "(Task Scheduler on Windows, launchd on macOS, cron on Linux). It:"
    )

    bullet(doc, " sleep stages (deep, light, REM, awake), overnight HRV, resting heart rate, "
        "body battery, steps, stress, and any workout details including HR zones and training effect.",
        bold_prefix="Pulls yesterday\u2019s data from Garmin Connect \u2014")

    bullet(doc, " Google Sheets serves as the human-readable source of truth. SQLite provides "
        "fast local queries and acts as a resilience layer \u2014 if the Sheets API is unavailable, "
        "data queues locally and retries on the next run. No data is ever lost.",
        bold_prefix="Writes to Google Sheets first, SQLite second.")

    bullet(doc, " Every numeric column gets color-graded (red-yellow-green) based on "
        "research-backed thresholds. Manual-entry columns are highlighted yellow. Weeks are visually "
        "grouped with alternating row bands.",
        bold_prefix="Applies visual formatting automatically.")

    bullet(doc, ", which computes the Readiness Score, generates insights, and optionally "
        "sends a morning notification.",
        bold_prefix="Triggers the analysis engine")

    body(doc,
        "For manual data entry (meals, workouts, subjective ratings), a voice-enabled Progressive Web App "
        "lets the user speak naturally \u2014 \u201chad two eggs with spinach at 7am\u201d \u2014 and "
        "Claude Haiku parses it into structured nutrition data, enriched by Nutritionix API lookups, "
        "and writes it directly to the correct spreadsheet tab."
    )

    add_divider(doc)

    # ── ANALYSIS ENGINE ──
    h1(doc, "The Analysis Engine")
    h2(doc, "Readiness Score: Methodology")

    body(doc,
        "The Readiness Score is a composite metric (1\u201310) built from four evidence-based components:"
    )

    t = doc.add_table(rows=5, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_cell_borders(t)
    t_data = [
        ["Component", "Weight", "Why This Weight"],
        ["HRV Status", "35%", "Strongest single predictor of next-day readiness (JAMA MESA 2020)"],
        ["Sleep Quality", "30%", "Second strongest; sleep architecture matters more than duration alone"],
        ["Resting Heart Rate", "20%", "Reliable but lagging indicator \u2014 downstream of HRV changes"],
        ["Subjective Wellness", "15%", "Self-reported energy + mood; downweighted when sleep debt detected (Van Dongen 2003)"],
    ]
    for r, row in enumerate(t_data):
        for c, val in enumerate(row):
            cell = t.rows[r].cells[c]
            if r == 0:
                set_cell_shading(cell, TABLE_HEADER_BG)
                run = cell.paragraphs[0].add_run(val)
                run.font.color.rgb = WHITE
                run.bold = True
            else:
                if r % 2 == 0:
                    set_cell_shading(cell, TABLE_ALT_BG)
                run = cell.paragraphs[0].add_run(val)
                run.font.color.rgb = DARK
                if c == 0:
                    run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
    style_table(t, col_widths=[1.5, 0.7, 4.3])

    body(doc,
        "Each component is scored against the individual\u2019s own rolling baseline \u2014 not "
        "population averages. An HRV of 38ms might be excellent for one person and concerning for another. "
        "The system learns your normal."
    )

    body_rich(doc, [
        ("Why sigmoid scoring instead of linear? ", True),
        ("Cognitive impairment doesn\u2019t scale linearly. Research shows it accelerates in the critical "
         "decision zone (the difference between \u201cFair\u201d and \u201cGood\u201d matters more than "
         "the difference between \u201cGood\u201d and \u201cOptimal\u201d). The sigmoid curve captures "
         "this \u2014 it\u2019s the same mathematical approach used by WHOOP and Oura in their "
         "proprietary scoring.", False),
    ])

    code_block(doc, "z = (today \u2212 30_day_mean) / 30_day_std_dev\nscore = 1 + 9 / (1 + e^(\u22121.5 \u00d7 z))")

    h2(doc, "Sleep Analysis")

    body(doc,
        "Rather than relying on Garmin\u2019s proprietary sleep score, the system computes an independent "
        "Sleep Analysis Score (0\u2013100) from seven metrics:"
    )

    bullet(doc, " (25 pts) \u2014 scored 0 at 4 hours or less, full marks at 7+", bold_prefix="Total sleep duration")
    bullet(doc, " (20 pts) \u2014 the restorative stage; research target is 15\u201320% of total sleep", bold_prefix="Deep sleep %")
    bullet(doc, " (20 pts) \u2014 critical for memory consolidation and emotional regulation", bold_prefix="REM sleep %")
    bullet(doc, " (15 pts) \u2014 parasympathetic recovery during sleep", bold_prefix="Overnight HRV")
    bullet(doc, " (10 pts) \u2014 sleep fragmentation degrades quality regardless of duration", bold_prefix="Awakenings")
    bullet(doc, " (10 pts) \u2014 Garmin\u2019s proprietary recovery metric, used as secondary signal", bold_prefix="Body battery gained")
    bullet(doc, " (+5 to \u221210 pts) \u2014 earlier bedtimes correlate with better sleep architecture", bold_prefix="Bedtime modifier")

    body(doc,
        "Sleep debt tracking follows the Van Dongen 2003 model from UPenn \u2014 the landmark study "
        "showing that subjective sleepiness plateaus after 3 days while cognitive decline continues "
        "silently. The system uses a 5-day weighted average (recent nights weighted more heavily) "
        "compared against a 30-day baseline to detect accumulating debt before the user feels it."
    )

    h2(doc, "Training Load")

    body(doc,
        "Training load analysis implements the Acute:Chronic Workload Ratio (ACWR) from Gabbett 2016 "
        "\u2014 the framework used by professional sports teams to manage injury risk:"
    )

    bullet(doc, " (0.8\u20131.3): training matches fitness level", bold_prefix="Sweet spot")
    bullet(doc, " (1.3\u20131.5): elevated risk, monitor recovery closely", bold_prefix="High zone")
    bullet(doc, " (>1.5): significantly elevated injury/illness risk", bold_prefix="Spike zone")
    bullet(doc, " (<0.8): insufficient stimulus", bold_prefix="Detraining")

    h2(doc, "Advanced Statistics (Built From Scratch)")

    body(doc,
        "Three analytical scripts provide deeper pattern analysis \u2014 all implemented using only "
        "Python\u2019s standard library (no NumPy, no SciPy):"
    )

    bullet(doc, " \u2014 Pearson correlations across all data domains with false discovery rate correction",
        bold_prefix="Correlation analysis")
    bullet(doc, " \u2014 OLS via the normal equation with variance inflation factor (VIF) for "
        "multicollinearity detection, plus leave-one-out cross-validation",
        bold_prefix="Multivariate regression")
    bullet(doc, " \u2014 answers questions like \u201cdoes a hard workout suppress HRV two days "
        "later?\u201d with autocorrelation-adjusted significance testing (Bayley & Hammersley 1946)",
        bold_prefix="Time-lagged correlation")

    body(doc,
        "Building these from scratch was a deliberate choice \u2014 it eliminates heavy dependencies, "
        "keeps the project portable, and demonstrates that the analytical depth doesn\u2019t require "
        "a data science stack."
    )

    add_divider(doc)

    # ── AI KNOWLEDGE SYSTEM ──
    h1(doc, "AI-Powered Knowledge System")
    h2(doc, "Four Claude Skills")

    body(doc,
        "The system extends Claude with four specialized skills that operate as domain-specific "
        "analytical tools:"
    )

    body_rich(doc, [
        ("/health-insight ", True),
        ("\u2014 Query-driven analysis that cross-references the research library with actual user data. "
         "Ask \u201cwhy is my HRV dropping?\u201d and it loads relevant domain knowledge, pulls your "
         "last 7\u201314 days of data, identifies the pattern, cites the research, and recommends "
         "specific actions.", False),
    ])

    body_rich(doc, [
        ("/update-intel ", True),
        ("\u2014 Processes new health research material (podcast transcripts, journal articles, book "
         "excerpts) into a structured knowledge base. Auto-classifies by domain, deduplicates against "
         "existing content, extracts quantifiable thresholds, and compiles findings into thematic "
         "Research Universe documents.", False),
    ])

    body_rich(doc, [
        ("/verify-intel ", True),
        ("\u2014 A fact-checking gate that evaluates health claims before they enter the library. "
         "Cross-references against a source hierarchy (meta-analyses > RCTs > clinical guidelines > "
         "expert consensus). Flags distortion patterns like cherry-picking, dose extrapolation, and "
         "animal-to-human leaps.", False),
    ])

    body_rich(doc, [
        ("/update-profile ", True),
        ("\u2014 Ingests personal medical documentation (lab results, diagnoses, provider notes) into "
         "a structured health profile. Conditions inform readiness weight adjustments, biomarkers "
         "correlate with physiological trends, and accommodations shape how recommendations are "
         "formatted and prioritized.", False),
    ])

    h2(doc, "Three-Layer Knowledge Hierarchy")

    t2 = doc.add_table(rows=4, cols=3)
    t2_data = [
        ["Layer", "What It Contains", "Purpose"],
        ["Research Universe Files\n(8 domains)", "Human-readable compilations with thematic organization\nand multi-source citations", "Deep reference for AI queries\nand human review"],
        ["Domain Briefs\n(<200 lines each)", "Token-efficient summaries: consensus positions,\nkey thresholds, open questions", "Fast context loading for\nreal-time analysis"],
        ["Runtime Knowledge JSON\n(100+ entries)", "Structured triggers with pattern matching,\ncognitive/energy impact, citations", "Auto-firing insights during\ndaily analysis"],
    ]
    for r, row in enumerate(t2_data):
        for c, val in enumerate(row):
            cell = t2.rows[r].cells[c]
            run = cell.paragraphs[0].add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
            if r == 0:
                set_cell_shading(cell, TABLE_HEADER_BG)
                run.font.color.rgb = WHITE
                run.bold = True
            else:
                run.font.color.rgb = DARK
                if r % 2 == 0:
                    set_cell_shading(cell, TABLE_ALT_BG)
    style_table(t2, col_widths=[2.0, 2.8, 1.7])

    body(doc,
        "New research can be added to the knowledge base and immediately influence daily analysis "
        "\u2014 no code changes required. The trigger system auto-fires insights when data patterns "
        "match research-backed thresholds."
    )

    add_divider(doc)

    # ── DELIVERY ──
    h1(doc, "Delivery: Raw Data to Phone")
    h2(doc, "Morning Briefing")

    body(doc, "Every morning, a push notification arrives via Pushover:")

    code_block(doc,
        "READINESS: 7.8 (Good) | Confidence: High\n\n"
        "SLEEP: 7.2h | Deep 18% | REM 21% | HRV 42ms | Bed 11:15pm\n"
        "  7d avg: Bed \u00b132min | Wake \u00b128min | Debt: 0.2h\n\n"
        "EXPECT: Above-baseline attention. Stable mood.\n\n"
        "FLAGS:\n"
        "  \u2022 HRV above baseline (z=+1.2) \u2014 autonomic ready\n"
        "  \u2022 ACWR 0.95 \u2014 sweet spot training load\n\n"
        "DO:\n"
        "  \u2022 Prioritize high-intensity or skill work\n"
        "  \u2022 Consider extra strength session if energy > 8/10"
    )

    body(doc,
        "All personal health information (condition names, medications, biomarker values) is "
        "automatically stripped before sending through the third-party notification service. "
        "PHI never leaves the local environment."
    )

    h2(doc, "Color-Graded Spreadsheets")

    bullet(doc, " deep sleep below 12% shows red (below clinical minimum); above 20% shows green (research target)",
        bold_prefix="Sleep metrics:")
    bullet(doc, " calibrated to individual baseline deviations, not population norms",
        bold_prefix="HRV:")
    bullet(doc, " discrete color bands \u2014 before 11 PM (green), 11 PM\u20131 AM (yellow), after 1 AM (red)",
        bold_prefix="Bedtime:")
    bullet(doc, " duration, distance, and calorie thresholds derived from ACSM guidelines",
        bold_prefix="Training metrics:")

    h2(doc, "Weekly Validation Loop")

    body(doc,
        "The system doesn\u2019t just predict \u2014 it validates. Every week, it automatically "
        "correlates its readiness predictions against actual next-day outcomes (self-reported morning "
        "energy and day rating). It computes Pearson correlations and flags when the model is drifting. "
        "This feedback loop is absent from most commercial platforms."
    )

    add_divider(doc)

    # ── ENGINEERING QUALITY ──
    h1(doc, "Engineering Quality")

    h3(doc, "Self-Verifying, Self-Healing")
    body(doc,
        "Every write to Google Sheets triggers two verification passes: structural verification "
        "(correct headers, valid data types, proper ordering) and formatting verification (conditional "
        "formatting rules exist AND numeric columns contain actual numbers, not text strings). If either "
        "check fails, the system auto-repairs and re-verifies. All repair logic is idempotent."
    )

    h3(doc, "Dual-Storage Resilience")
    body(doc,
        "Google Sheets + SQLite with a retry queue. SQLite writes always succeed locally. Failed Sheets "
        "writes queue in pending_sync.json and retry on the next run. No data is ever lost."
    )

    h3(doc, "Cross-Platform Portability")
    body(doc,
        "Runs identically on Windows, macOS, and Linux. Credentials use the OS-native keyring "
        "(Windows Credential Manager / macOS Keychain / libsecret). All file paths are dynamic. "
        "Migration to a new machine is an 8-step checklist."
    )

    h3(doc, "Zero Credential Exposure")
    body(doc,
        "Garmin password lives only in the OS keyring. Google service account key is gitignored. "
        "Pushover tokens are in .env (gitignored). Protected health information stays in a gitignored "
        "profiles/ directory. Notifications are sanitized. Work logs and commit messages never contain "
        "medical details."
    )

    h3(doc, "Batch Write Optimization")
    body(doc,
        "Google Sheets enforces a 60-request-per-minute quota. The system reads entire columns, "
        "modifies in memory, and writes back in a single batch call. Mixed data types are handled with "
        "a split strategy: RAW mode for dates/times, USER_ENTERED for numeric columns."
    )

    add_divider(doc)

    # ── BUILT WITH CLAUDE ──
    h1(doc, "Built With Claude")

    h3(doc, "1. Architect")
    body(doc,
        "Claude designed the system architecture from the ground up \u2014 the 4-layer pipeline, "
        "the dual-storage strategy, the sigmoid scoring methodology, the knowledge hierarchy, and "
        "the PHI boundary model. Every architectural decision was discussed, debated, and justified "
        "before implementation."
    )

    h3(doc, "2. Engineer")
    body(doc,
        "Claude wrote every line of Python in the project: 15+ production scripts with full error "
        "handling, custom implementations of OLS regression, Pearson correlation, and lag analysis "
        "without external math libraries, a serverless voice logger PWA with TOTP authentication, "
        "a database migration system, cross-platform scheduler setup, and comprehensive verification "
        "and self-repair systems."
    )

    h3(doc, "3. Runtime Analytical Engine")
    body(doc,
        "Claude operates as the analytical brain at runtime through the skills system. It loads domain "
        "knowledge, pulls real data from Google Sheets, computes trends and anomalies, cross-references "
        "against the research library, and synthesizes personalized answers with citations. This isn\u2019t "
        "a chatbot answering health questions \u2014 it\u2019s an AI analyst with access to actual "
        "biometric data and a curated research library."
    )

    add_divider(doc)

    # ── TECH STACK ──
    h1(doc, "Technology Stack")

    ts = doc.add_table(rows=13, cols=3)
    ts_data = [
        ["Layer", "Technology", "Purpose"],
        ["Data Source", "Garmin Connect API", "50+ daily biometrics"],
        ["Data Source", "Nutritionix API", "Nutrition lookup for voice logger"],
        ["Storage (Primary)", "Google Sheets API (gspread)", "Human-readable source of truth"],
        ["Storage (Backup)", "SQLite", "Local mirror, offline queries, resilience"],
        ["Analysis", "Python 3.14 (stdlib only)", "Readiness scoring, correlation, regression"],
        ["Knowledge", "Claude AI (skills system)", "Research curation, insight generation"],
        ["NLP", "Claude Haiku", "Voice-to-structured-data parsing"],
        ["Notifications", "Pushover API", "Morning briefing delivery"],
        ["Voice Input", "Web Speech API + Vercel", "PWA for nutrition/workout logging"],
        ["Auth", "TOTP (HMAC-based)", "Voice logger authentication"],
        ["Credentials", "OS keyring (cross-platform)", "Zero-file credential storage"],
        ["Scheduling", "Task Scheduler / launchd / cron", "Nightly automated sync"],
    ]
    for r, row in enumerate(ts_data):
        for c, val in enumerate(row):
            cell = ts.rows[r].cells[c]
            run = cell.paragraphs[0].add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
            if r == 0:
                set_cell_shading(cell, TABLE_HEADER_BG)
                run.font.color.rgb = WHITE
                run.bold = True
            else:
                run.font.color.rgb = DARK
                if r % 2 == 0:
                    set_cell_shading(cell, TABLE_ALT_BG)
    style_table(ts, col_widths=[1.5, 2.5, 2.5])

    add_divider(doc)

    # ── WHAT MAKES THIS IMPRESSIVE ──
    h1(doc, "What Makes This Impressive")

    body_rich(doc, [
        ("It\u2019s not a dashboard \u2014 it\u2019s an analyst. ", True),
        ("Most health platforms show you charts. This one tells you what the charts mean, why they "
         "matter today, and what to do about it \u2014 backed by cited research, calibrated to your "
         "personal baseline, and aware of your medical context.", False),
    ])

    body_rich(doc, [
        ("The knowledge system learns without code changes. ", True),
        ("New research is ingested, fact-checked, compiled into thematic documents, and converted into "
         "auto-firing triggers \u2014 all through the AI skill system. The analytical engine gets "
         "smarter over time without touching a single line of Python.", False),
    ])

    body_rich(doc, [
        ("It validates its own predictions. ", True),
        ("The weekly correlation between readiness scores and actual next-day outcomes creates a "
         "feedback loop that most commercial platforms lack entirely. The system knows when it\u2019s "
         "wrong.", False),
    ])

    body_rich(doc, [
        ("The statistics are built from scratch. ", True),
        ("Pearson correlation with FDR correction, OLS regression with VIF multicollinearity detection, "
         "time-lagged correlation with autocorrelation-adjusted significance testing \u2014 all in pure "
         "Python. No NumPy, no SciPy. Portable and lightweight while delivering graduate-level "
         "statistical analysis.", False),
    ])

    body_rich(doc, [
        ("PHI handling is enterprise-grade. ", True),
        ("Medical data stays in a gitignored directory. Notifications are sanitized. Work logs are "
         "generic. Commit messages reveal nothing. Designed as if it would be audited \u2014 because "
         "personal health data deserves that standard.", False),
    ])

    body_rich(doc, [
        ("Every write is verified. ", True),
        ("Every single write triggers structural and formatting verification, with auto-repair if "
         "anything fails. The user has never had to manually fix a spreadsheet formatting issue. "
         "That level of operational reliability is unusual in any project.", False),
    ])

    # ── CLOSING ──
    add_divider(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run(
        "15,000+ lines of production Python  \u00b7  8 research domains  \u00b7  "
        "100+ knowledge entries  \u00b7  50+ daily biometrics"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Architected, engineered, and powered by Claude.")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED
    run.italic = True

    # Save
    out = os.path.join(DIR, "Health Tracker - Executive Brief.docx")
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
